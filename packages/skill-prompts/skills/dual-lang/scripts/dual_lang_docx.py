#!/usr/bin/env python3
"""
dual_lang_docx.py — Builds a dual-language (VI+EN) .docx from a Vietnamese source.

Each top-level body paragraph is followed immediately by its English translation.
Each paragraph inside a table cell is followed by its English translation within
the same cell.

English paragraphs are formatted as italic + blue to visually distinguish them from
the Vietnamese source text (matching the LS Law Firm bilingual template style).

Output file uses the -V,E.docx suffix unless --output is provided.

Workflow:
  1. --extract-only  : dump all text segments as JSON to stdout
  2. Build translation map JSON: { "original vi text": "english translation", ... }
  3. --translation-map <json> : produce the dual-language .docx
"""

import argparse
import copy
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)

# Hex color applied to all English translation paragraphs (blue, matching bilingual template)
EN_COLOR_HEX = "2E75B6"


def _top_level_body_paras(doc):
    """Only paragraphs that are direct children of the document body (not inside tables)."""
    body = doc.element.body
    return [p for p in doc.paragraphs if p._element.getparent() is body]


def collect_segments(doc):
    segments = []

    for para in _top_level_body_paras(doc):
        text = para.text.strip()
        if text:
            segments.append({"type": "paragraph", "text": text})

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if text:
                        segments.append({
                            "type": "cell_paragraph",
                            "table_index": ti,
                            "row_index": ri,
                            "col_index": ci,
                            "para_index": pi,
                            "text": text,
                        })

    return segments


def _make_en_para_xml(source_para, en_text):
    """Clone source paragraph XML (preserving pPr/alignment/spacing), apply EN formatting.

    EN paragraphs are italic + blue (EN_COLOR_HEX). Font name and size are inherited
    from the first run of the source paragraph so spacing stays consistent.
    """
    xml = copy.deepcopy(source_para._element)

    # Strip all runs and hyperlinks — we'll add a single clean run
    for tag in (qn('w:r'), qn('w:hyperlink'), qn('w:ins'), qn('w:del')):
        for elem in xml.findall(tag):
            xml.remove(elem)

    # Build rPr: inherit font name/size from source, override with italic + blue
    rpr = OxmlElement('w:rPr')

    if source_para.runs:
        src_rpr = source_para.runs[0]._element.find(qn('w:rPr'))
        if src_rpr is not None:
            # Carry over font face and size only — not bold/color/italic from source
            for tag in (qn('w:rFonts'), qn('w:sz'), qn('w:szCs'), qn('w:lang')):
                elem = src_rpr.find(tag)
                if elem is not None:
                    rpr.append(copy.deepcopy(elem))

    # Italic (w:i for Latin, w:iCs for complex scripts like Vietnamese)
    rpr.append(OxmlElement('w:i'))
    rpr.append(OxmlElement('w:iCs'))

    # Blue color
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), EN_COLOR_HEX)
    rpr.append(color_elem)

    r_elem = OxmlElement('w:r')
    r_elem.append(rpr)

    t_elem = OxmlElement('w:t')
    t_elem.text = en_text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)
    xml.append(r_elem)
    return xml


def apply_dual_language(doc, tmap):
    """
    Insert EN translations into the document in-place.

    Body paragraphs: EN paragraph inserted immediately after its VI paragraph.
    Cell paragraphs: EN paragraph inserted immediately after its VI paragraph in the cell.

    Both loops run in reverse order so earlier insertions don't shift subsequent positions.
    """
    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in reversed(list(cell.paragraphs)):
                    text = para.text.strip()
                    if text and text in tmap:
                        en_text = tmap[text]
                        if en_text and en_text != text:
                            en_xml = _make_en_para_xml(para, en_text)
                            para._element.addnext(en_xml)

    # Top-level body paragraphs
    for para in reversed(_top_level_body_paras(doc)):
        text = para.text.strip()
        if text and text in tmap:
            en_text = tmap[text]
            if en_text and en_text != text:
                en_xml = _make_en_para_xml(para, en_text)
                para._element.addnext(en_xml)


def main():
    parser = argparse.ArgumentParser(description="Dual-language (VI+EN) .docx builder")
    parser.add_argument("input", help="Path to the source Vietnamese .docx file")
    parser.add_argument("--extract-only", action="store_true",
                        help="Print all text segments as JSON and exit (no output file)")
    parser.add_argument("--translation-map",
                        help="Path to JSON file mapping original VI text to EN translation")
    parser.add_argument("--output", help="Output path. Default: <stem>-V,E.docx beside the input")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(input_path))
    segments = collect_segments(doc)

    if args.extract_only:
        print(json.dumps(segments, ensure_ascii=False, indent=2))
        return

    if not args.translation_map:
        print("ERROR: Provide --translation-map <json> or use --extract-only.", file=sys.stderr)
        print("\nExtracted segments:")
        print(json.dumps(segments, ensure_ascii=False, indent=2))
        sys.exit(1)

    map_path = Path(args.translation_map)
    if not map_path.exists():
        print(f"ERROR: Translation map not found: {map_path}", file=sys.stderr)
        sys.exit(1)

    with open(map_path, encoding="utf-8") as f:
        tmap = json.load(f)

    apply_dual_language(doc, tmap)

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.parent / f"{input_path.stem}-V,E{input_path.suffix}"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    translated = sum(1 for s in segments if s.get("text", "") in tmap)
    print(f"Saved: {output_path}")
    print(f"Segments translated: {translated} / {len(segments)}")

    missing = [s["text"] for s in segments if s.get("text", "") not in tmap]
    if missing:
        print(f"\nWARNING: {len(missing)} segment(s) had no translation in the map:")
        for t in missing[:5]:
            print(f"  - {t[:80]}")
        if len(missing) > 5:
            print(f"  ... and {len(missing) - 5} more.")


if __name__ == "__main__":
    main()
