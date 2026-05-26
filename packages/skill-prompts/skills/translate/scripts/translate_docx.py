#!/usr/bin/env python3
"""
translate_docx.py — Format-preserving .docx translation scaffold for vn-legal-translate skill.

This script extracts text from a .docx file paragraph-by-paragraph and table cell-by-cell,
then writes a new .docx where the translated text replaces the original while run-level
formatting (font name, size, bold, italic, underline, color) is preserved.

IMPORTANT: This script does NOT call any external translation API. It prints each text
segment to stdout so Claude (the caller) can provide the translations interactively,
OR it accepts a pre-built translation map via --translation-map <json-file>.

Typical workflow when called by Claude:
  1. Run with --extract-only to dump all text segments to a JSON file.
  2. Claude translates each segment and saves the results to a translation map JSON.
  3. Run again with --translation-map <file> to produce the translated .docx.

Usage:
  python3 translate_docx.py <input.docx> --direction <vi-to-en|en-to-vi> [options]

Options:
  --output <path>           Output .docx path. Default: <input>-<lang>.docx
  --extract-only            Print all text segments as JSON and exit (no output file).
  --translation-map <json>  Path to JSON file mapping original text → translated text.
  --direction <vi-to-en|en-to-vi>   Translation direction (required).
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


def collect_segments(doc: Document) -> list[dict]:
    """Walk the document and collect all text segments with their location."""
    segments = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            segments.append({
                "type": "paragraph",
                "index": i,
                "text": text,
                "translated": None,
            })

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if text:
                        segments.append({
                            "type": "table_cell",
                            "table_index": ti,
                            "row_index": ri,
                            "col_index": ci,
                            "para_index": pi,
                            "text": text,
                            "translated": None,
                        })

    return segments


def apply_translations(doc: Document, translation_map: dict[str, str]) -> int:
    """Replace text in each run while preserving run-level formatting. Returns count of replacements."""
    replaced = 0

    def translate_paragraph(para):
        nonlocal replaced
        original_full = para.text.strip()
        if not original_full or original_full not in translation_map:
            return
        translated_full = translation_map[original_full]
        if not translated_full or translated_full == original_full:
            return

        # Strategy: put all translated text into the first non-empty run,
        # clear the rest. This preserves the first run's formatting.
        runs = [r for r in para.runs if r.text.strip()]
        if not runs:
            return

        first_run = runs[0]
        first_run.text = translated_full
        for run in runs[1:]:
            run.text = ""
        replaced += 1

    for para in doc.paragraphs:
        translate_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    translate_paragraph(para)

    return replaced


def main():
    parser = argparse.ArgumentParser(description="Format-preserving .docx translation scaffold")
    parser.add_argument("input", help="Path to the source .docx file")
    parser.add_argument("--direction", choices=["vi-to-en", "en-to-vi"], required=True)
    parser.add_argument("--output", help="Output .docx path")
    parser.add_argument("--extract-only", action="store_true",
                        help="Print text segments as JSON and exit")
    parser.add_argument("--translation-map", help="JSON file: {original_text: translated_text}")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(input_path))
    segments = collect_segments(doc)

    if args.extract_only:
        print(json.dumps(segments, ensure_ascii=False, indent=2))
        return

    if not args.translation_map:
        print("ERROR: Provide --translation-map <json> or use --extract-only.", file=sys.stderr)
        print("\nExtracted segments (translate these, then rerun with --translation-map):")
        print(json.dumps(segments, ensure_ascii=False, indent=2))
        sys.exit(1)

    map_path = Path(args.translation_map)
    if not map_path.exists():
        print(f"ERROR: Translation map not found: {map_path}", file=sys.stderr)
        sys.exit(1)

    with open(map_path, encoding="utf-8") as f:
        translation_map = json.load(f)

    replaced_count = apply_translations(doc, translation_map)

    lang_suffix = "en" if args.direction == "vi-to-en" else "vi"
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.parent / f"{input_path.stem}-{lang_suffix}{input_path.suffix}"

    doc.save(str(output_path))
    print(f"Saved translated document to: {output_path}")
    print(f"Paragraphs/cells translated: {replaced_count} / {len(segments)} segments")

    untranslated = [s["text"] for s in segments if s["text"] not in translation_map]
    if untranslated:
        print(f"\nWARNING: {len(untranslated)} segment(s) had no translation in the map:")
        for t in untranslated[:10]:
            print(f"  - {t[:80]}")
        if len(untranslated) > 10:
            print(f"  ... and {len(untranslated) - 10} more.")


if __name__ == "__main__":
    main()
